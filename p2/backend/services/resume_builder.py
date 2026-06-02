import io
import json
import re
from collections import Counter
from typing import Dict, List, Optional

from docx import Document
from services.advanced_resume_analyzer import AdvancedResumeAnalyzer


class ResumeBuilderService:
    @staticmethod
    def _normalize_text(value: Optional[str]) -> str:
        return value.strip() if isinstance(value, str) else ''

    @staticmethod
    def _normalize_items(items: Optional[List[dict]]) -> List[dict]:
        if not isinstance(items, list):
            return []
        normalized = []
        for item in items:
            if not isinstance(item, dict):
                continue
            normalized.append({
                'role': ResumeBuilderService._normalize_text(item.get('role')),
                'company': ResumeBuilderService._normalize_text(item.get('company')),
                'dates': ResumeBuilderService._normalize_text(item.get('dates')),
                'details': ResumeBuilderService._normalize_text(item.get('details')),
                'school': ResumeBuilderService._normalize_text(item.get('school')),
                'degree': ResumeBuilderService._normalize_text(item.get('degree')),
            })
        return normalized

    @staticmethod
    def normalize_resume_data(data: dict) -> dict:
        return {
            'title': ResumeBuilderService._normalize_text(data.get('title')),
            'template_id': ResumeBuilderService._normalize_text(data.get('template_id')) or 'modern',
            'contact': {
                'name': ResumeBuilderService._normalize_text(data.get('contact', {}).get('name')),
                'email': ResumeBuilderService._normalize_text(data.get('contact', {}).get('email')),
                'phone': ResumeBuilderService._normalize_text(data.get('contact', {}).get('phone')),
                'website': ResumeBuilderService._normalize_text(data.get('contact', {}).get('website')),
                'location': ResumeBuilderService._normalize_text(data.get('contact', {}).get('location')),
            },
            'summary': ResumeBuilderService._normalize_text(data.get('summary')),
            'experience': ResumeBuilderService._normalize_items(data.get('experience')),
            'education': ResumeBuilderService._normalize_items(data.get('education')),
            'skills': ResumeBuilderService._normalize_text(data.get('skills')),
            'projects': ResumeBuilderService._normalize_text(data.get('projects')),
            'certifications': ResumeBuilderService._normalize_text(data.get('certifications')),
            'job_description': ResumeBuilderService._normalize_text(data.get('job_description')),
        }

    @staticmethod
    def format_resume_text(data: dict) -> str:
        normalized = ResumeBuilderService.normalize_resume_data(data)
        lines = []

        contact = normalized['contact']
        if contact['name']:
            lines.append(contact['name'])
        contact_items = [contact['email'], contact['phone'], contact['website']]
        contact_line = ' | '.join([item for item in contact_items if item])
        if contact_line:
            lines.append(contact_line)
        if contact['location']:
            lines.append(contact['location'])

        if normalized['summary']:
            lines.append('\nSUMMARY')
            lines.append(normalized['summary'])

        if normalized['experience']:
            lines.append('\nEXPERIENCE')
            for item in normalized['experience']:
                title = ' '.join([item['role'], '—', item['company']]).strip(' —')
                lines.append(title or 'Role — Company')
                if item['dates']:
                    lines.append(item['dates'])
                if item['details']:
                    lines.append(item['details'])
                lines.append('')

        if normalized['education']:
            lines.append('\nEDUCATION')
            for item in normalized['education']:
                title = ' '.join([item['degree'], '—', item['school']]).strip(' —')
                lines.append(title or 'Degree — School')
                if item['dates']:
                    lines.append(item['dates'])
                if item['details']:
                    lines.append(item['details'])
                lines.append('')

        if normalized['skills']:
            lines.append('\nSKILLS')
            lines.append(normalized['skills'])

        if normalized['projects']:
            lines.append('\nPROJECTS')
            lines.append(normalized['projects'])

        if normalized['certifications']:
            lines.append('\nCERTIFICATIONS')
            lines.append(normalized['certifications'])

        return '\n'.join(line for line in lines if line is not None)

    @staticmethod
    def format_resume_markdown(data: dict) -> str:
        normalized = ResumeBuilderService.normalize_resume_data(data)
        blocks = []
        contact = normalized['contact']

        if contact['name']:
            blocks.append(f"# {contact['name']}")
        contact_items = [contact['email'], contact['phone'], contact['website']]
        if any(contact_items):
            blocks.append(' | '.join([item for item in contact_items if item]))
        if contact['location']:
            blocks.append(contact['location'])

        if normalized['summary']:
            blocks.append('## Summary')
            blocks.append(normalized['summary'])

        if normalized['experience']:
            blocks.append('## Experience')
            for item in normalized['experience']:
                blocks.append(f"### {item['role'] or 'Role'} — {item['company'] or 'Company'}")
                if item['dates']:
                    blocks.append(f"*{item['dates']}*")
                if item['details']:
                    blocks.append(item['details'])

        if normalized['education']:
            blocks.append('## Education')
            for item in normalized['education']:
                blocks.append(f"### {item['degree'] or 'Degree'} — {item['school'] or 'School'}")
                if item['dates']:
                    blocks.append(f"*{item['dates']}*")
                if item['details']:
                    blocks.append(item['details'])

        if normalized['skills']:
            blocks.append('## Skills')
            blocks.append(normalized['skills'])

        if normalized['projects']:
            blocks.append('## Projects')
            blocks.append(normalized['projects'])

        if normalized['certifications']:
            blocks.append('## Certifications')
            blocks.append(normalized['certifications'])

        return '\n\n'.join(blocks)

    @staticmethod
    def format_resume_html(data: dict, template_id: str = 'modern') -> str:
        normalized = ResumeBuilderService.normalize_resume_data(data)
        sections = []
        contact = normalized['contact']
        theme_class = 'resume-template-modern'
        if template_id == 'professional':
            theme_class = 'resume-template-professional'
        elif template_id == 'minimal':
            theme_class = 'resume-template-minimal'

        header_lines = []
        if contact['name']:
            header_lines.append(f"<h1>{contact['name']}</h1>")
        contact_items = [contact['email'], contact['phone'], contact['website']]
        if any(contact_items):
            header_lines.append(f"<p>{' | '.join([item for item in contact_items if item])}</p>")
        if contact['location']:
            header_lines.append(f"<p>{contact['location']}</p>")

        if normalized['summary']:
            sections.append(f"<section><h2>Summary</h2><p>{normalized['summary']}</p></section>")

        if normalized['experience']:
            content = ''.join([
                '<div class="resume-block">'
                f"<h3>{item['role'] or 'Role'} — {item['company'] or 'Company'}</h3>"
                f"<p class='small'>{item['dates'] or ''}</p>"
                f"<p>{item['details'] or ''}</p>"
                '</div>'
                for item in normalized['experience']
            ])
            sections.append(f"<section><h2>Experience</h2>{content}</section>")

        if normalized['education']:
            content = ''.join([
                '<div class="resume-block">'
                f"<h3>{item['degree'] or 'Degree'} — {item['school'] or 'School'}</h3>"
                f"<p class='small'>{item['dates'] or ''}</p>"
                f"<p>{item['details'] or ''}</p>"
                '</div>'
                for item in normalized['education']
            ])
            sections.append(f"<section><h2>Education</h2>{content}</section>")

        if normalized['skills']:
            sections.append(f"<section><h2>Skills</h2><p>{normalized['skills']}</p></section>")

        if normalized['projects']:
            sections.append(f"<section><h2>Projects</h2><p>{normalized['projects']}</p></section>")

        if normalized['certifications']:
            sections.append(f"<section><h2>Certifications</h2><p>{normalized['certifications']}</p></section>")

        body = '<div class="resume-wrapper">' + ''.join(sections) + '</div>'
        html = f"<html><head><meta charset='utf-8'/><style>body{{font-family:Arial,Helvetica,sans-serif;color:#111;background:#fff;padding:24px;}}h1{{font-size:26px;margin-bottom:8px;}}h2{{font-size:16px;text-transform:uppercase;letter-spacing:.08em;margin:20px 0 8px;color:#0f766e;}}h3{{font-size:14px;margin:10px 0 2px;}}p{{margin:0 0 10px;line-height:1.6;}}.small{{font-size:12px;color:#4b5563;}}.resume-wrapper{{max-width:760px;}}.resume-block{{margin-bottom:12px;}}</style></head><body>{''.join(header_lines)}{body}</body></html>"
        return html

    @staticmethod
    def format_resume_pdf(data: dict) -> bytes:
        normalized = ResumeBuilderService.normalize_resume_data(data)
        lines = ResumeBuilderService.format_resume_text(normalized).splitlines()
        if not lines:
            lines = ['']

        def escape_pdf_text(text: str) -> str:
            return text.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)').replace('\r', '')

        pages = []
        current_lines = []
        max_lines_per_page = 48
        for line in lines:
            current_lines.append(line)
            if len(current_lines) >= max_lines_per_page:
                pages.append(current_lines)
                current_lines = []
        if current_lines:
            pages.append(current_lines)

        pdf_objects = []
        page_kids = []
        font_obj_id = 3 + len(pages) * 2

        for index, page_lines in enumerate(pages):
            page_id = 3 + index * 2
            content_id = page_id + 1
            page_kids.append(f"{page_id} 0 R")

            content_lines = ["BT", "/F1 12 Tf", "72 760 Td"]
            for line in page_lines:
                content_lines.append(f"({escape_pdf_text(line)}) Tj")
                content_lines.append("T*")
            content_lines.append("ET")
            content_stream = "\n".join(content_lines).encode('latin-1', 'replace')

            pdf_objects.append((content_id, f"{content_id} 0 obj << /Length {len(content_stream)} >>\nstream\n".encode('latin-1') + content_stream + b"\nendstream\n"))
            page_obj = (
                f"{page_id} 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                f"/Contents {content_id} 0 R /Resources << /Font << /F1 {font_obj_id} 0 R >> >> >> endobj\n"
            )
            pdf_objects.insert(index, (page_id, page_obj.encode('latin-1')))

        pages_obj = (
            f"2 0 obj << /Type /Pages /Kids [{' '.join(page_kids)}] /Count {len(pages)} >> endobj\n"
        ).encode('latin-1')
        catalog_obj = b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        font_obj = (
            f"{font_obj_id} 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        ).encode('latin-1')

        body_parts = [catalog_obj, pages_obj] + [obj for _, obj in pdf_objects] + [font_obj]
        offsets = []
        cursor = 0
        for part in body_parts:
            offsets.append(cursor)
            cursor += len(part)

        xref = [b"xref\n0 %d\n0000000000 65535 f \n" % (len(body_parts) + 1)]
        for offset in offsets:
            xref.append(f"{offset:010} 00000 n \n".encode('latin-1'))

        trailer = (
            f"trailer << /Size {len(body_parts) + 1} /Root 1 0 R >>\nstartxref\n{cursor}\n%%EOF\n"
        ).encode('latin-1')

        return b"%PDF-1.4\n" + b"".join(body_parts) + b"".join(xref) + trailer

    @staticmethod
    def format_resume_word(data: dict) -> bytes:
        normalized = ResumeBuilderService.normalize_resume_data(data)
        document = Document()
        contact = normalized['contact']
        if contact['name']:
            document.add_heading(contact['name'], level=0)
        contact_line = ' | '.join([item for item in [contact['email'], contact['phone'], contact['website']] if item])
        if contact_line:
            document.add_paragraph(contact_line)
        if contact['location']:
            document.add_paragraph(contact['location'])

        if normalized['summary']:
            document.add_heading('Summary', level=1)
            document.add_paragraph(normalized['summary'])

        if normalized['experience']:
            document.add_heading('Experience', level=1)
            for item in normalized['experience']:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{item['role'] or 'Role'} — {item['company'] or 'Company'}").bold = True
                if item['dates']:
                    paragraph.add_run(f"\n{item['dates']}")
                if item['details']:
                    paragraph.add_run(f"\n{item['details']}")

        if normalized['education']:
            document.add_heading('Education', level=1)
            for item in normalized['education']:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{item['degree'] or 'Degree'} — {item['school'] or 'School'}").bold = True
                if item['dates']:
                    paragraph.add_run(f"\n{item['dates']}")
                if item['details']:
                    paragraph.add_run(f"\n{item['details']}")

        if normalized['skills']:
            document.add_heading('Skills', level=1)
            document.add_paragraph(normalized['skills'])

        if normalized['projects']:
            document.add_heading('Projects', level=1)
            document.add_paragraph(normalized['projects'])

        if normalized['certifications']:
            document.add_heading('Certifications', level=1)
            document.add_paragraph(normalized['certifications'])

        stream = io.BytesIO()
        document.save(stream)
        return stream.getvalue()

    @staticmethod
    def extract_skill_tags(skills_text: str) -> List[str]:
        if not skills_text:
            return []
        raw_tags = re.split(r'[\n,;]+', skills_text)
        return [tag.strip() for tag in raw_tags if tag.strip()]

    @staticmethod
    def extract_keywords(text: str) -> List[str]:
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        filtered = [word for word in words if word not in {'and', 'the', 'for', 'with', 'that', 'from', 'this', 'will', 'have', 'has', 'had', 'but', 'not'}]
        frequencies = Counter(filtered)
        top = [word for word, _ in frequencies.most_common(40)]
        return top

    @staticmethod
    def match_job_description(resume_text: str, job_description: str) -> Dict:
        resume_tokens = set(ResumeBuilderService.extract_keywords(resume_text))
        jd_tokens = set(ResumeBuilderService.extract_keywords(job_description))
        matched = sorted(token for token in jd_tokens if token in resume_tokens)
        missing = sorted(token for token in jd_tokens if token not in resume_tokens)
        score = 0
        if jd_tokens:
            score = int(len(matched) / len(jd_tokens) * 100)
        return {
            'match_percentage': score,
            'matched_keywords': matched[:40],
            'missing_keywords': missing[:40],
            'resume_keywords': sorted(resume_tokens)[:40],
            'job_description_keywords': sorted(jd_tokens)[:40],
        }

    @staticmethod
    def render_preview(data: dict) -> Dict:
        normalized = ResumeBuilderService.normalize_resume_data(data)
        return {
            'title': normalized['title'] or 'Untitled Resume',
            'template': normalized['template_id'],
            'text': ResumeBuilderService.format_resume_text(normalized),
            'markdown': ResumeBuilderService.format_resume_markdown(normalized),
            'html': ResumeBuilderService.format_resume_html(normalized, normalized['template_id']),
            'skills': ResumeBuilderService.extract_skill_tags(normalized['skills']),
        }
